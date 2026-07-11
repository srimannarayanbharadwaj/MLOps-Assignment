"""Generate the final submission Word document for the MLOps assignment."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "report" / "FINAL_REPORT.docx"


def add_title_page(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Heart Disease MLOps\nFinal Report")
    run.bold = True
    run.font.size = Pt(20)

    doc.add_paragraph()
    details = [
        "Student: Sriman Narayan Bharadwaj R",
        "Course: MLOps Assignment 01 (2026)",
        "Repository: https://github.com/srimannarayanbharadwaj/MLOps-Assignment",
        "Date: 12-07-2026",
    ]
    for line in details:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_paragraph(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def add_image(doc: Document, relative_path: str, caption: str, width: float = 6.0) -> None:
    image_path = ROOT / relative_path
    if image_path.exists():
        doc.add_picture(str(image_path), width=Inches(width))
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].italic = True
    else:
        doc.add_paragraph(f"[Missing image: {relative_path}]")


def add_metrics_table(doc: Document) -> None:
    table = doc.add_table(rows=3, cols=5)
    table.style = "Table Grid"
    headers = ["Model", "Accuracy", "Precision", "Recall", "ROC-AUC"]
    rows = [
        ["Logistic Regression", "0.8447", "0.8475", "0.8050", "0.9188"],
        ["Random Forest", "0.8416", "0.8388", "0.8127", "0.9163"],
    ]
    for col, header in enumerate(headers):
        table.rows[0].cells[col].text = header
    for row_idx, row_data in enumerate(rows, start=1):
        for col, value in enumerate(row_data):
            table.rows[row_idx].cells[col].text = value


def build_document() -> Document:
    doc = Document()
    add_title_page(doc)

    add_heading(doc, "1. Executive Summary")
    add_paragraph(
        doc,
        "This project builds an end-to-end MLOps pipeline for the UCI Cleveland "
        "Heart Disease dataset. The workflow covers data ingestion, exploratory "
        "analysis, model training, experiment tracking, packaging, automated testing, "
        "containerized serving, Kubernetes deployment, and basic monitoring. Two "
        "models were compared using stratified 5-fold cross-validation: Logistic "
        "Regression and Random Forest. Logistic Regression was selected for "
        "deployment because it achieved the best ROC-AUC (0.9188) while remaining "
        "simple to explain and efficient to serve in production.",
    )

    add_heading(doc, "2. Project Setup", level=1)
    add_heading(doc, "2.1 Environment", level=2)
    add_bullets(
        doc,
        [
            "Python version: 3.12 (used because MLflow UI failed on Python 3.14).",
            "Virtual environment: .venv312",
            "Core libraries: pandas, scikit-learn, mlflow, fastapi, uvicorn, pytest, prometheus-client",
        ],
    )

    add_heading(doc, "2.2 Repository Structure", level=2)
    add_code(
        doc,
        "heart-disease-mlops/\n"
        "  data/raw, data/processed\n"
        "  src/               # download, cleaning, EDA, training, API\n"
        "  tests/             # unit tests\n"
        "  models/            # packaged pipeline\n"
        "  .github/workflows/ # CI pipeline\n"
        "  k8s/               # deployment manifests\n"
        "  screenshots/       # evidence from each phase\n"
        "  report/            # final report and generated metrics",
    )

    add_heading(doc, "2.3 Reproduction Commands", level=2)
    add_code(
        doc,
        "python -m venv .venv312\n"
        ".\\.venv312\\Scripts\\Activate.ps1\n"
        "pip install -r requirements.txt\n"
        "python -m src.data_download\n"
        "python -m src.eda\n"
        "python -m src.train\n"
        "python -m src.package_model",
    )

    add_heading(doc, "3. Phase 1 — Data and EDA", level=1)
    add_heading(doc, "3.1 Dataset", level=2)
    add_bullets(
        doc,
        [
            "Source: UCI Heart Disease dataset, Cleveland subset",
            "Rows after cleaning: 303",
            "Features: 13 clinical attributes plus binary target",
            "Missing values appeared as '?' in ca and thal and were handled in preprocessing",
            "Target rule: 0 = no disease, values 1+ collapsed to 1 = disease",
        ],
    )

    add_heading(doc, "3.2 Class Balance", level=2)
    add_bullets(doc, ["No disease (0): 164", "Disease (1): 139"])
    add_paragraph(
        doc,
        "The classes are moderately imbalanced. Accuracy alone can be misleading "
        "because a model could appear strong by predicting the majority class more "
        "often. For this medical screening task, recall, precision, and ROC-AUC are "
        "more informative because they reflect how well the model catches actual "
        "disease cases and ranks risk correctly.",
    )
    add_image(doc, "screenshots/phase1_eda/class_balance.png", "Figure 1: Class balance")

    add_heading(doc, "3.3 Feature Distributions", level=2)
    add_paragraph(
        doc,
        "The numeric histograms show that age is roughly bell-shaped, while "
        "cholesterol and blood pressure have wider spreads with some extreme values. "
        "Oldpeak and thalach also vary noticeably across patients. Because the "
        "numeric features use different scales, StandardScaler was applied before "
        "model training.",
    )
    add_image(
        doc,
        "screenshots/phase1_eda/numeric_feature_histograms.png",
        "Figure 2: Numeric feature histograms",
    )

    add_heading(doc, "3.4 Correlation Heatmap", level=2)
    add_paragraph(
        doc,
        "The heatmap shows moderate relationships between several clinical variables "
        "and the target. Features such as cp, thalach, and oldpeak stand out as "
        "useful predictors. Some input variables are also correlated with each other, "
        "which is expected in medical data. Tree models can tolerate this naturally, "
        "while Logistic Regression still performed well after scaling and one-hot encoding.",
    )
    add_image(
        doc,
        "screenshots/phase1_eda/correlation_heatmap.png",
        "Figure 3: Correlation heatmap",
    )

    add_heading(doc, "4. Phase 2 — Feature Engineering and Modeling", level=1)
    add_heading(doc, "4.1 Preprocessing Pipeline", level=2)
    add_bullets(
        doc,
        [
            "Numeric features: age, trestbps, chol, thalach, oldpeak -> StandardScaler",
            "Categorical features: sex, cp, fbs, restecg, exang, slope, ca, thal -> OneHotEncoder",
            "Combined with ColumnTransformer inside a sklearn Pipeline",
        ],
    )

    add_heading(doc, "4.2 Cross-Validation Results", level=2)
    add_metrics_table(doc)
    add_paragraph(doc, "Evaluation used StratifiedKFold with 5 folds.")

    add_heading(doc, "4.3 Model Selection Justification", level=2)
    add_paragraph(
        doc,
        "Logistic Regression was selected as the production model because it achieved "
        "the highest ROC-AUC (0.9188), indicating stronger overall ranking ability "
        "across thresholds. Random Forest had slightly better recall (0.8127 vs 0.8050), "
        "which matters clinically because missed disease cases can be more serious than "
        "false alarms. However, the recall difference is small, and Logistic Regression "
        "is easier to interpret and faster to deploy in an API service.",
    )

    add_heading(doc, "5. Phase 3 — Experiment Tracking (MLflow)", level=1)
    add_paragraph(
        doc,
        "Each training run was logged in MLflow with model parameters, cross-validation "
        "metrics, confusion matrix images, CSV summaries, and the full sklearn pipeline. "
        "MLflow made it easy to compare Logistic Regression and Random Forest side by side "
        "without manually opening multiple output files.",
    )
    add_image(
        doc,
        "screenshots/phase3_mlflow/01_experiments_page.png",
        "Figure 4: MLflow experiments page",
        width=6.2,
    )
    add_image(
        doc,
        "screenshots/phase3_mlflow/02_run_training_table.png",
        "Figure 5: MLflow training run comparison",
        width=6.2,
    )
    add_image(
        doc,
        "screenshots/phase3_mlflow/04_logistic_artifacts.png",
        "Figure 6: Logistic Regression run artifacts",
        width=6.2,
    )

    add_heading(doc, "6. Phase 4 — Packaging", level=1)
    add_bullets(
        doc,
        [
            "models/heart_disease_pipeline.joblib — preprocessing + model pipeline",
            "models/model_metadata.json — feature lists, training rows, CV metrics",
            "models/sample_input.json — sample inference payload",
            "requirements.txt — pinned dependencies for reproducibility",
        ],
    )
    add_paragraph(
        doc,
        "A fresh clone can reproduce the packaged artifact by installing requirements "
        "and running python -m src.package_model.",
    )

    add_heading(doc, "7. Phase 5 — Unit Tests and CI/CD", level=1)
    add_bullets(
        doc,
        [
            "tests/test_data.py validates null handling and target binarization",
            "tests/test_model.py checks prediction shape and probability validity",
            "tests/test_api.py checks /health and /metrics endpoints",
        ],
    )
    add_paragraph(
        doc,
        "GitHub Actions runs lint, pytest, and the training/package workflow in order. "
        "This prevents broken cleaning logic or model code from being merged silently.",
    )
    add_image(
        doc,
        "screenshots/Screenshot 2026-07-11 214818.png",
        "Figure 7: Successful GitHub Actions workflow",
        width=6.2,
    )

    add_heading(doc, "8. Phase 6 — Containerization (Docker + FastAPI)", level=1)
    add_bullets(
        doc,
        [
            "GET /health — service and model status",
            "POST /predict — prediction, confidence, and risk label",
            "GET /metrics — Prometheus metrics",
        ],
    )
    add_image(doc, "screenshots/docker_run.png", "Figure 8: Docker container running")
    add_image(
        doc,
        "screenshots/phase6_docker/03_curl_predict_success.png",
        "Figure 9: Successful Docker /predict response",
    )
    add_paragraph(
        doc,
        'Sample response: {"prediction": 0, "confidence": 0.8292, '
        '"risk_label": "no_disease", "model_name": "logistic_regression"}',
    )

    add_heading(doc, "9. Phase 7 — Kubernetes Deployment (Minikube)", level=1)
    add_paragraph(
        doc,
        "The Docker image was deployed to a local Minikube cluster using Kubernetes "
        "Deployment and Service manifests. Readiness and liveness probes use /health.",
    )
    add_code(
        doc,
        "minikube start\n"
        "docker build -t heart-disease-api:latest .\n"
        "minikube image load heart-disease-api:latest\n"
        "kubectl apply -f k8s/\n"
        "kubectl port-forward service/heart-disease-api 8080:8000",
    )
    add_image(
        doc,
        "screenshots/phase7_k8s/minicube start success.png",
        "Figure 10: Minikube started successfully",
    )
    add_image(
        doc,
        "screenshots/phase7_k8s/k8s deployment.png",
        "Figure 11: Kubernetes deployment and service created",
    )
    add_image(
        doc,
        "screenshots/phase7_k8s/k8s port-forward.png",
        "Figure 12: Port-forward to the Kubernetes service",
    )
    add_image(
        doc,
        "screenshots/phase7_k8s/k8s health success.png",
        "Figure 13: Successful /health response from Kubernetes deployment",
    )

    add_heading(doc, "10. Phase 8 — Monitoring", level=1)
    add_paragraph(
        doc,
        "The FastAPI service includes request logging middleware and a Prometheus "
        "/metrics endpoint. Logs capture timestamp, endpoint, method, status code, "
        "and latency. In production, these signals would help diagnose slow endpoints "
        "or rising error rates.",
    )
    add_image(
        doc,
        "screenshots/phase8_monitoring/Request logs.png",
        "Figure 14: Request logging output",
        width=6.2,
    )
    add_image(
        doc,
        "screenshots/phase8_monitoring/metrics endpoint.png",
        "Figure 15: Prometheus /metrics endpoint",
        width=6.2,
    )

    add_heading(doc, "11. System Architecture", level=1)
    add_paragraph(
        doc,
        "Pipeline flow: UCI CSV -> data_download.py -> cleaned dataset -> EDA and "
        "train.py -> MLflow tracking and generated metrics -> package_model.py -> "
        "joblib pipeline -> FastAPI -> Docker image -> Minikube deployment -> "
        "/predict, /health, and /metrics endpoints. GitHub Actions and pytest validate "
        "the pipeline automatically on each push.",
    )
    add_code(
        doc,
        "UCI CSV\n"
        "  -> src/data_download.py\n"
        "  -> data/processed/heart_cleveland_clean.csv\n"
        "  -> src/eda.py + src/train.py\n"
        "  -> MLflow + report/generated metrics\n"
        "  -> src/package_model.py\n"
        "  -> models/heart_disease_pipeline.joblib\n"
        "  -> src/api.py (FastAPI)\n"
        "  -> Docker image\n"
        "  -> Kubernetes deployment (Minikube)\n"
        "  -> /predict, /health, /metrics\n\n"
        "GitHub Actions + pytest validate each stage automatically.",
    )

    add_heading(doc, "12. Challenges and Lessons Learned", level=1)
    add_bullets(
        doc,
        [
            "Python 3.14 caused MLflow UI failures; Python 3.12 resolved the issue.",
            "Docker and Minikube port conflicts required alternate ports or stopping old tunnels.",
            "Pinned requirements improved reproducibility across machines.",
            "MLflow significantly simplified experiment comparison and artifact management.",
        ],
    )

    add_heading(doc, "13. Conclusion", level=1)
    add_paragraph(
        doc,
        "This assignment delivered a complete MLOps workflow from raw medical data to "
        "a monitored production-style API. The final Logistic Regression pipeline "
        "achieved strong cross-validated performance and was packaged, tested, "
        "containerized, and deployed on Kubernetes with logging and metrics support.",
    )

    add_heading(doc, "Appendix A — Repository Link", level=1)
    add_paragraph(doc, "https://github.com/srimannarayanbharadwaj/MLOps-Assignment")

    add_heading(doc, "Appendix B — Video Demonstration", level=1)
    add_paragraph(
        doc,
        "Add the recorded end-to-end walkthrough video link or filename here before final submission.",
    )

    return doc


def main() -> None:
    doc = build_document()
    doc.save(OUTPUT)
    print(f"Created {OUTPUT}")


if __name__ == "__main__":
    main()
